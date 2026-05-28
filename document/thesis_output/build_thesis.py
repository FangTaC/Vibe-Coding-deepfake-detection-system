from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "thesis_source.md"
OUTPUT = ROOT / "智能体协作的深度伪造人脸检测系统_毕业论文初稿.docx"


def parse_source(path: Path) -> tuple[dict[str, str], list[str]]:
    text = path.read_text(encoding="utf-8")
    meta_text, body_text = text.split("---\n", 1)
    meta: dict[str, str] = {}
    for line in meta_text.splitlines():
        if not line.strip():
            continue
        key, value = line.split(":", 1)
        meta[key.strip()] = value.strip()
    return meta, body_text.splitlines()


def set_run_font(run, size_pt: float, bold: bool = False, italic: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def set_paragraph_format(paragraph, first_line_chars: int = 2, line_spacing: float = 1.5) -> None:
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(10.5 * first_line_chars)
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def add_text_with_citations(paragraph, text: str, size_pt: float = 12.0, bold: bool = False) -> None:
    parts = re.split(r"(\[\d+(?:[\-,， ]\d+)*\])", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        set_run_font(run, size_pt=size_pt, bold=bold)
        if re.fullmatch(r"\[\d+(?:[\-,， ]\d+)*\]", part):
            run.font.superscript = True


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size_pt=10.5)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "目录将在 Word 中更新域后生成"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(text)
    run._r.append(fld_end)
    set_run_font(run, size_pt=12.0)


def set_page_number_format(section, fmt: str | None = None, start: int | None = None) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    if fmt:
        pg_num_type.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num_type.set(qn("w:start"), str(start))


def add_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "double")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "auto")
    borders.append(bottom)


def configure_section(section) -> None:
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.0)


def add_cover(doc: Document, meta: dict[str, str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    run = p.add_run(meta["SCHOOL"])
    set_run_font(run, size_pt=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run("本科毕业论文（设计）")
    set_run_font(run, size_pt=22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(55)
    run = p.add_run(meta["TITLE_CN"])
    set_run_font(run, size_pt=18, bold=True)

    for label, key in [
        ("学院", "COLLEGE"),
        ("专业", "MAJOR"),
        ("学生姓名", "STUDENT"),
        ("学号", "STUDENT_ID"),
        ("指导教师", "SUPERVISOR"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if label == "学院":
            p.paragraph_format.space_before = Pt(70)
        run = p.add_run(f"{label}：{meta[key]}")
        set_run_font(run, size_pt=14, bold=False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    run = p.add_run(meta["DATE_CN"])
    set_run_font(run, size_pt=14)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(180)
    run = p.add_run("毕业论文（设计）原创性声明和版权使用授权书")
    set_run_font(run, size_pt=16, bold=True)
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_chars=2)
    add_text_with_citations(
        p,
        "说明：本页通常按学院统一模板另页签署。当前自动生成文档仅保留占位说明，提交定稿前请替换为学院正式模板。",
    )


def add_front_section(doc: Document, meta: dict[str, str], body_blocks: list[tuple[str, list[str]]]) -> None:
    front = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(front)
    front.header.is_linked_to_previous = False
    front.footer.is_linked_to_previous = False
    set_page_number_format(front, fmt="upperRoman", start=1)
    fp = front.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)

    abstract_cn = dict(body_blocks).get("中文摘要", [])
    abstract_en = dict(body_blocks).get("Abstract", [])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta["TITLE_CN"])
    set_run_font(run, size_pt=15, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("摘  要")
    set_run_font(run, size_pt=16, bold=True)
    for line in abstract_cn:
        if not line.strip():
            continue
        p = doc.add_paragraph()
        if line.startswith("关键词："):
            add_text_with_citations(p, line, size_pt=12, bold=True)
            p.paragraph_format.space_before = Pt(12)
        else:
            set_paragraph_format(p, first_line_chars=2)
            add_text_with_citations(p, line)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta["TITLE_EN"])
    set_run_font(run, size_pt=15, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Abstract")
    set_run_font(run, size_pt=16, bold=True)
    for line in abstract_en:
        if not line.strip():
            continue
        p = doc.add_paragraph()
        if line.startswith("Keywords:"):
            add_text_with_citations(p, line, size_pt=12, bold=True)
            p.paragraph_format.space_before = Pt(12)
        else:
            set_paragraph_format(p, first_line_chars=0)
            add_text_with_citations(p, line)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("目  录")
    set_run_font(run, size_pt=16, bold=True)
    p = doc.add_paragraph()
    add_toc(p)


def add_body_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    set_page_number_format(section, fmt="decimal", start=1)

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run("湖南大学本科毕业论文")
    set_run_font(run, size_pt=10.5)
    add_bottom_border(hp)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)


def build_blocks(lines: list[str]) -> list[tuple[str, list[str]]]:
    blocks: list[tuple[str, list[str]]] = []
    current_title = ""
    current_lines: list[str] = []
    for line in lines:
        if line.startswith("# "):
            if current_title:
                blocks.append((current_title, current_lines))
            current_title = line[2:].strip()
            current_lines = []
        else:
            current_lines.append(line)
    if current_title:
        blocks.append((current_title, current_lines))
    return blocks


def add_body(doc: Document, blocks: list[tuple[str, list[str]]]) -> None:
    in_body = False
    for title, lines in blocks:
        if title in {"中文摘要", "Abstract", "目录"}:
            continue
        if not in_body:
            add_body_section(doc)
            in_body = True
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if title.startswith("第") and "章" in title else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title)
        set_run_font(run, size_pt=16 if title.startswith("第") and "章" in title else 14, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)

        for raw in lines:
            line = raw.strip()
            if not line:
                continue
            if line.startswith("## "):
                p = doc.add_paragraph()
                run = p.add_run(line[3:].strip())
                set_run_font(run, size_pt=12, bold=True)
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(8)
                continue
            if line.startswith("IMAGE:"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line.replace("IMAGE:", "【图片占位】", 1).strip())
                set_run_font(run, size_pt=10.5, italic=True)
                continue
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line_chars=2)
            add_text_with_citations(p, line)


def main() -> None:
    meta, lines = parse_source(SOURCE)
    blocks = build_blocks(lines)
    doc = Document()
    configure_section(doc.sections[0])
    doc.sections[0].different_first_page_header_footer = True
    add_cover(doc, meta)
    add_front_section(doc, meta, blocks)
    add_body(doc, blocks)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
